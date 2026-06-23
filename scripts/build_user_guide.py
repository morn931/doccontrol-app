"""
Build the CoreDocs User Guide (Word .docx) AND the placeholder screenshots the in-app
Guide reads from /public/guide.

Run:  python scripts/build_user_guide.py   (from the doccontrol-app folder)
Output:
  - public/guide/*.png      (placeholder slots — the in-app Guide + the .docx use these)
  - CoreDocs-User-Guide.docx

Screenshots: the PNGs are dashed "replace me" placeholders. In copilot mode, open each
live CoreDocs page and save a real screenshot over public/guide/<slug>.png, then re-run
this script to refresh the .docx with the real images.

KEEP THE CONTENT BELOW IN SYNC WITH lib/guide/registry.ts (same screens, same tips).
"""
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(ROOT, "public", "guide")
os.makedirs(ASSETS, exist_ok=True)

BASE_URL = "https://docs.coreflow.build"   # live CoreDocs (alias of doccontrol-app.vercel.app)

# ── Screens (mirror of lib/guide/registry.ts) ─────────────────────────────────
# (slug, route, title, intro, [tips])
SCREENS = [
    ("dashboard", "/dashboard", "Dashboard",
     "Your home screen — an at-a-glance view of the document-control workload and the way in to every part of CoreDocs.",
     ["The left sidebar is your main menu — what you see depends on your role (Document Controller, Reviewer, Engineering/Project Manager, Vendor or Admin).",
      "The summary cards show the current state of play — batches awaiting action, reviews due, documents by status. Click one to jump straight in.",
      "Use the top-right menu for your account and to sign out. The Guide button (top of every page) explains the screen you are on."]),
    ("batches", "/batches", "Incoming Batches",
     "Vendor document batches as they arrive — the Document Controller's inbox for logging and distributing new submissions.",
     ["Each batch is a set of documents submitted together (from Aconex / the vendor). Open one to see its documents and metadata.",
      "Register the batch into the master register, then assign the documents to the right reviewers / disciplines.",
      "Track a batch's progress from received -> under review -> returned, so nothing sits unactioned."]),
    ("reviews", "/reviews", "My Reviews",
     "The documents assigned to you to review, and where you record your review outcome. The list is split into Pending / In Progress and Completed, with an overdue banner so nothing slips.",
     ["Work the top of the list first - anything flagged overdue holds up the whole document cycle. Open a document to view it (and any markups) in the Review Chain.",
      "Record a formal outcome code: A1 (approved), B1/B2 (approved with comments), C1 / D1 (revise & resubmit), Q1 (for quotation), V1 / S1 (information / superseded). Add your comments alongside the code.",
      "Your outcome and comments flow back to the Document Controller and onto the transmittal - once set, the document moves to Completed."]),
    ("transmittals", "/transmittals", "Transmittals",
     "The Transmittal Register - formal issue and receipt of documents, the auditable record of what was sent to whom, when, and why.",
     ["Create a transmittal to issue documents (e.g. returning reviewed vendor docs, or sending for construction) with a cover sheet.",
      "Each transmittal has a unique number and lists its documents, revisions and the reason for issue.",
      "Open a past transmittal to see its full history - the permanent record for audits and claims."]),
    ("documents", "/documents", "Document Search",
     "Search the full document register (3,500+ documents) - find any deliverable by number, title, package, vendor or status.",
     ["Use Smart Search to describe the document in plain language ('the HV single-line diagram for the substation') - it matches on meaning, not just exact text.",
      "Narrow the list with the Package / Vendor / Source / Sector filters and the Awarded / Unawarded toggle.",
      "Open a document to see its full history - every revision, review and transmittal it has been through. This is the quickest way to answer 'what's the latest revision / status of X?'."]),
    ("mddr", "/mddr", "MDDR - Master Register",
     "The Master Document & Drawing Register - the controlled master list combining the SDDR, CDDL and MDDR into one register of every deliverable (6,000+ entries) with its current revision, status and progress.",
     ["Each row is a deliverable with its document number, title, package/discipline, current revision and status. Use Columns to choose what's shown.",
      "Run Sync Progress to refresh deliverable progress from the latest data; use Upload Register to bulk-load or update entries, and Export CSV for an offline copy or the client return.",
      "This is the single source of truth for deliverable progress - it drives the Reporting dashboards."]),
    ("reporting", "/reporting", "Reporting",
     "Progress and status reports built live from the register - for the project team and the client.",
     ["The reports are: Progress Dashboard (overall % complete), Engineering Tracker, Package Progress Summary, PPE Phase 1 Engineering Deliverables, and the P6 Activity-ID Progress Export.",
      "Use the P6 Activity-ID export to feed deliverable progress straight back into the Primavera P6 schedule.",
      "Everything reads live from the MDDR, so the numbers are always current - no manual spreadsheet upkeep."]),
    ("admin-import", "/admin/import", "Import & Sync",
     "Bring SharePoint data into CoreDocs - the automatic SharePoint sync and a manual CSV import (admin only). Always run a dry run first to preview before committing.",
     ["Automatic SharePoint Sync pulls the Approver Picks and Document Approval lists straight from SharePoint via Microsoft Graph (no CSV needed) - it runs every day at 02:00 UTC. Use 'Sync now (force update)' to refresh immediately, 'Sync changes only' for just the deltas, or 'Preview (dry run)' to see what would change.",
      "To import from a CSV instead: pick the Import Source (e.g. Approver Picks - Batch records), choose a mode - Dry Run (validate only), Full (insert/update all) or Incremental (new records only) - then upload the CSV.",
      "Always start with a Dry Run: it validates and shows what would happen without changing anything. Re-running imports is safe; check the result summary after each run."]),
    ("admin-users", "/admin/users", "Users",
     "Manage who can access the Document Control platform - user accounts and roles (admin only).",
     ["Each user has a role that controls their menu and permissions: Admin, Reviewer, Document Controller, Engineering/Project Manager or Vendor. The badge next to each name shows their current role.",
      "Use 'Add User' to invite someone, or 'Edit' to change a person's role - it takes effect on their next page load.",
      "Keep the reviewer list current so incoming batches can be assigned to the right people."]),
    ("admin-vendors", "/admin/vendors", "Vendors & Packages",
     "The project packages and the vendor each is awarded to (admin only). PPE's own engineering scope sits under package K124.",
     ["Each row is a package (e.g. E101 - 36MVA High Speed Diesel Generator) with an 'Awarded: <vendor>' badge, or 'Not awarded yet' if the contract isn't placed.",
      "Set the awarded vendor as packages are placed - this is what lets batches, the register and reporting group documents by package and vendor correctly."]),
]

SECTIONS = [
    ("Getting started", ["dashboard"]),
    ("Document-control workflow", ["batches", "reviews", "transmittals"]),
    ("Register, search & reporting", ["documents", "mddr", "reporting"]),
    ("Admin", ["admin-import", "admin-users", "admin-vendors"]),
]

BY_SLUG = {s[0]: s for s in SCREENS}

# ── Placeholder screenshot generator ──────────────────────────────────────────
def placeholder(slug, title, url):
    p = os.path.join(ASSETS, slug + ".png")
    if os.path.exists(p) and os.path.getsize(p) > 30000:
        return p  # a real screenshot is already in place — don't overwrite it
    W, H = 1600, 900
    img = Image.new("RGB", (W, H), (245, 246, 248))
    d = ImageDraw.Draw(img)
    for x in range(0, W, 26):
        d.line([(x, 4), (x + 14, 4)], fill=(170, 175, 182), width=3)
        d.line([(x, H - 4), (x + 14, H - 4)], fill=(170, 175, 182), width=3)
    for y in range(0, H, 26):
        d.line([(4, y), (4, y + 14)], fill=(170, 175, 182), width=3)
        d.line([(W - 4, y), (W - 4, y + 14)], fill=(170, 175, 182), width=3)
    try:
        f_big = ImageFont.truetype("arialbd.ttf", 54)
        f_med = ImageFont.truetype("arial.ttf", 38)
        f_small = ImageFont.truetype("arial.ttf", 30)
    except Exception:
        f_big = f_med = f_small = ImageFont.load_default()
    def center(text, y, font, fill):
        w = d.textlength(text, font=font)
        d.text(((W - w) / 2, y), text, font=font, fill=fill)
    center("SCREENSHOT", 300, f_big, (120, 126, 134))
    center(title, 400, f_med, (90, 96, 104))
    center(url, 470, f_small, (140, 146, 154))
    center("Replace this image with a screenshot of the page above.", 560, f_small, (160, 165, 172))
    img.save(p)
    return p

# ── Doc helpers ───────────────────────────────────────────────────────────────
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t):
    para = doc.add_paragraph(); para.add_run(t); return para
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def shot(slug, title, url):
    img = placeholder(slug, title, url)
    doc.add_picture(img, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Screen: " + title); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Title ─────────────────────────────────────────────────────────────────────
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("CoreDocs"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = RGBColor(0x16, 0x3A, 0x5F)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Vendor Document Approval, Transmittals & the Master Register (MDDR)"); r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x55, 0x5B, 0x63)
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run("User Guide  ·  PPE Technologies  ·  Coreflow")
sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub3.add_run("Generated from the in-app guide — June 2026"); r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
doc.add_page_break()

h1("What this guide covers")
for heading, slugs in SECTIONS:
    bullet(heading + ": " + ", ".join(BY_SLUG[s][2] for s in slugs))
doc.add_page_break()

# ── Sections ──────────────────────────────────────────────────────────────────
for heading, slugs in SECTIONS:
    h1(heading)
    for sg in slugs:
        slug, route, title, intro, tips = BY_SLUG[sg]
        h2(title)
        p(intro)
        shot(slug, title, BASE_URL + route)
        # Reviews has a second screenshot: the review-detail page (Review Chain + outcome codes).
        if slug == "reviews" and os.path.exists(os.path.join(ASSETS, "review-detail.png")):
            doc.add_picture(os.path.join(ASSETS, "review-detail.png"), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run("Screen: My Reviews - review detail (Review Chain & outcome codes)")
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        for t in tips:
            bullet(t)

doc.add_paragraph()
end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("End of guide — CoreDocs, PPE Technologies"); r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out = os.path.join(ROOT, "CoreDocs-User-Guide.docx")
doc.save(out)
print("Saved:", out)
print("Placeholders in:", ASSETS)
